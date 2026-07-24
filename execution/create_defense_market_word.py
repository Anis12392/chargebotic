"""
Chargebotic — Full Defense Market Research Report
Generates a complete Word (.docx) document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Color palette ─────────────────────────────────────────────────────────────
C_BLACK      = RGBColor(0x0A, 0x0A, 0x0A)
C_DARK       = RGBColor(0x1A, 0x1A, 0x2E)
C_AMBER      = RGBColor(0xD4, 0x7E, 0x0F)
C_RED        = RGBColor(0xC0, 0x39, 0x2B)
C_GREEN      = RGBColor(0x1A, 0x6B, 0x3A)
C_BLUE       = RGBColor(0x1A, 0x4A, 0x7A)
C_GRAY       = RGBColor(0x55, 0x55, 0x55)
C_LIGHT      = RGBColor(0xF5, 0xF5, 0xF5)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_TABLE_HEAD = RGBColor(0x1A, 0x2E, 0x1A)
C_TABLE_ALT  = RGBColor(0xF0, 0xF4, 0xF0)


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_border_bottom(para, color="C8C8C8", size=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    set_para_border_bottom(p, color="D4810F", size=8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = C_AMBER
    run.font.name = "Calibri"
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_DARK
    run.font.name = "Calibri"
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_BLUE
    run.font.name = "Calibri"
    return p


def body(text, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_BLACK
    run.font.name = "Calibri"
    return p


def bullet(text, bold_part=None, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = "Calibri"
        r1.font.color.rgb = color or C_AMBER
        rest = text[len(bold_part):]
        if rest:
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
            r2.font.name = "Calibri"
            r2.font.color.rgb = C_BLACK
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = C_BLACK
    return p


def callout(text, bg=C_LIGHT, border_color="D4810F"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), border_color)
        pBdr.append(el)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.italic = True
    run.font.color.rgb = C_BLACK
    return p


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TABLE_HEAD)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        run.font.name = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_i, row_data in enumerate(rows):
        row = t.rows[r_i + 1]
        alt = (r_i % 2 == 1)
        for c_i, cell_text in enumerate(row_data):
            cell = row.cells[c_i]
            if alt:
                set_cell_bg(cell, C_TABLE_ALT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = C_BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CHARGEBOTIC")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = C_DARK
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Defense Market Research — Full Report")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = C_AMBER
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Tethered Powerline Drone · Battlefield Energy Delivery")
run.font.size = Pt(12)
run.font.color.rgb = C_GRAY
run.font.name = "Calibri"
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
run = p.add_run("June 2026  ·  Confidential  ·  Defense Segment Only")
run.font.size = Pt(10)
run.font.color.rgb = C_GRAY
run.font.name = "Calibri"

set_para_border_bottom(p, color="D4810F", size=12)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("KEY FIGURES AT A GLANCE")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = C_AMBER
run.font.name = "Calibri"

table(
    ["Metric", "Figure", "Context"],
    [
        ["DoD Annual Energy Spend",      "$20B+/year",       "Fuel + electricity across all installations"],
        ["Fully Burdened Fuel Cost",     "$400/gallon",      "At a remote forward operating base (vs. $3 pump)"],
        ["Military Power Market (TAM)",  "$7.7B–$10.8B",     "2024; growing at 8.5–19% CAGR"],
        ["Chargebotic SAM",              "$400M–$800M",      "Deployable off-grid tactical power sub-segment"],
        ["US Army BCTs",                 "59 total",         "32 active-duty + 27 National Guard"],
        ["USSOCOM Personnel",            "73,000 SOF",       "Across 80+ countries simultaneously"],
        ["Convoy Casualty Rate",         "1 per 24 convoys", "Afghanistan fuel resupply data"],
        ["Direct Competitors",           "0",                "No equivalent product exists today"],
        ["Top B2B2G Partner",            "Chariot Defense",  "$41M funded by a16z — distributes power, needs our source"],
    ],
    col_widths=[2.5, 1.8, 3.5]
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
h1("1. The Problem — Why the Military Needs This")

body(
    "The United States military has a well-documented and critical energy logistics problem. "
    "Diesel generators dominate power at forward operating bases (FOBs). Every gallon of diesel "
    "requires a supply convoy. Every convoy is a target. The cost — financial and human — is enormous."
)

h2("1.1 The Financial Cost of Military Energy")

bullet("$20B+ per year", bold_part="$20B+ per year", color=C_RED)
body("    The US Department of Defense is the single largest institutional energy consumer on the planet, spending over $20 billion per year on fuel and electricity to power everything from fighter jets to forward operating bases.", space_after=Pt(2))

bullet("$400 per gallon (fully burdened)", bold_part="$400 per gallon (fully burdened)", color=C_RED)
body("    Fuel that costs ~$3/gallon at the pump can reach $400/gallon when fully burdened — accounting for transportation through hostile or remote terrain (e.g., convoyed through Pakistan, over the Hindu Kush, or across the Amu Darya river to forward bases).", space_after=Pt(2))

bullet("$20.2B/year on air conditioning alone", bold_part="$20.2B/year on air conditioning alone", color=C_RED)
body("    The US military spent an estimated $20.2 billion annually on air conditioning structures in Iraq and Afghanistan — one of the largest single energy line items at any FOB.", space_after=Pt(2))

bullet("68M gallons/month in theater", bold_part="68M gallons/month in theater", color=C_AMBER)
body("    In 2008, DoD supplied an average of 68 million gallons of fuel per month to support US military forces in Iraq and Afghanistan. A single large Army division can consume up to 6,000 gallons of fuel per day.", space_after=Pt(2))

h2("1.2 The Human Cost of Fuel Logistics")

bullet("1 casualty per 24 fuel convoys", bold_part="1 casualty per 24 fuel convoys", color=C_RED)
body("    The casualty factor for fuel resupply in Afghanistan was 0.042 — meaning nearly one casualty for every 24 fuel resupply convoys.", space_after=Pt(2))

bullet("50%+ of convoy casualties tied to fuel", bold_part="50%+ of convoy casualties tied to fuel", color=C_RED)
body("    Attacks on fuel convoys accounted for more than 50% of casualties in Afghanistan and 30% in Iraq during 2009. Fuel is described by analysts as America's biggest battlefield vulnerability.", space_after=Pt(2))

bullet("70–80% of convoy weight is fuel and water", bold_part="70–80% of convoy weight is fuel and water", color=C_AMBER)
body("    This weight crowds out ammunition, food, and medical supplies. Energy logistics directly degrades combat effectiveness.", space_after=Pt(2))

h2("1.3 Why Existing Solutions Are Insufficient")

body(
    "Current solutions reduce fuel consumption but do not eliminate the logistics problem. "
    "Hybrid solar-battery systems reduce generator diesel consumption by 30–50% in optimal conditions — "
    "but remain weather-dependent, require significant setup time, and still require diesel as backup. "
    "Chargebotic eliminates the fuel dependency entirely for any location near existing power line infrastructure, "
    "delivering continuous, weather-independent, infrastructure-tapped energy."
)

callout(
    "KEY INSIGHT:  Power lines are already deployed near every human settlement, road, military base, and training facility. "
    "Chargebotic is the first system that taps this existing infrastructure autonomously — no fuel, no weather dependency, no convoy."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MARKET SIZE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Market Size")

body(
    "Multiple independent research firms have sized the global military power market between $7.7B and $10.8B in 2024, "
    "with projections ranging to $13.6B–$24.9B by 2030–2032. The variance reflects different scope definitions. "
    "All sources agree on double-digit CAGR driven by electrification of the battlefield, energy independence mandates, "
    "and increasing drone/C4ISR power demand."
)

h2("2.1 Market Segments and Sizing")

table(
    ["Segment", "2024 Value", "Projection", "CAGR", "Source"],
    [
        ["Military Power Solutions (broad TAM)", "$7.7B–$10.8B", "$13.6B–$24.9B by 2032", "8.5–9.7%", "Mordor Intelligence / IMARC"],
        ["Next-Gen Military Power Supply",       "$6.4B",        "$9.5B by 2032",          "5.1%",     "SkyQuestT / GM Insights"],
        ["Military Microgrid Market",            "$2.0B",        "High growth",             "19.1%",    "GM Insights 2024"],
        ["Military DC Microgrid",                "$700M",        "—",                       "High",     "GM Insights 2024"],
        ["Power Generator for Military",         "$1.2B",        "$1.7B by 2035",           "3.8%",     "Future Market Insights"],
        ["CHARGEBOTIC SAM (deployable power)",   "$400M–$800M",  "$1.5B+ by 2030",          "~15%",     "Chargebotic estimate"],
    ],
    col_widths=[2.4, 1.4, 1.8, 1.0, 1.7]
)

h2("2.2 Why the Market Is Growing Fast")

bullet("Electrification of the battlefield", bold_part="Electrification of the battlefield", color=C_BLUE)
body("    Drones, autonomous systems, electronic warfare, directed energy weapons, and C4ISR all require growing amounts of electrical power at the tactical edge.", space_after=Pt(2))

bullet("DoD mandate to reduce fuel dependency", bold_part="DoD mandate to reduce fuel dependency", color=C_BLUE)
body("    The Army Operational Energy Strategy formally mandates reducing FOB fuel consumption. Congressional and GAO pressure following convoy casualty data has created top-down budget support.", space_after=Pt(2))

bullet("DARPA active investment in energy innovation", bold_part="DARPA active investment in energy innovation", color=C_BLUE)
body("    DARPA's POWER program (Persistent Optical Wireless Energy Relay, $10M+), ExPEDitions battery program (targeting 5–10× energy density), and multiple SBIR cycles signal sustained institutional demand.", space_after=Pt(2))

bullet("Tethered drones becoming formal Army programs (2025)", bold_part="Tethered drones becoming formal Army programs (2025)", color=C_BLUE)
body("    The Army is converting tethered drone experimentation into formal program-of-record status — directly opening the procurement pathway for Chargebotic's hardware.", space_after=Pt(2))

bullet("US share dominates", bold_part="US share dominates", color=C_BLUE)
body("    The United States accounts for approximately 86% of the North American military power market, driven by extensive modernization programs and the largest defense budget in the world.", space_after=Pt(2))

h2("2.3 Chargebotic's Serviceable Market")

body(
    "Chargebotic's SAM is the deployable off-grid power sub-segment — bases, training sites, and field units that are "
    "(a) near power line infrastructure and (b) currently paying the high fully-burdened cost of diesel logistics. "
    "This is estimated conservatively at $400M–$800M in 2024, growing to $1.5B+ by 2030 at a ~15% CAGR. "
    "The US alone hosts 750+ overseas military installations in 80+ countries, the majority of which have power lines within operational radius."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CUSTOMERS (DIRECT)
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Direct Customers — Who Buys, How Many, What They Need")

body(
    "Chargebotic has three tiers of direct government customers: the US Army (largest buyer by volume), "
    "USSOCOM (highest priority, fastest procurement), and the US Marine Corps (secondary, expeditionary). "
    "Allied forces and NATO partners represent an international expansion layer."
)

# ── 3.1 US ARMY
h2("3.1 US Army — Primary Customer")

h3("Who They Are")
body(
    "The US Army is the world's largest land force and the dominant consumer of forward-deployed power generation "
    "equipment. Army Brigade Combat Teams (BCTs) are the primary operational unit and the relevant buying unit "
    "for FOB power systems."
)

h3("Key Numbers")
table(
    ["Metric", "Figure", "Notes"],
    [
        ["Active-Duty BCTs",          "32",              "14 Infantry, 11 Armored, 7 Stryker"],
        ["National Guard BCTs",       "27",              "20 Infantry, 5 Armored, 2 Stryker"],
        ["Total BCTs",                "59",              "Each runs its own FOB power infrastructure"],
        ["Soldiers per BCT",          "4,000–4,700",     "Commanded by a Colonel"],
        ["US overseas installations", "750+",            "In 80+ countries"],
        ["Permanently stationed abroad","173,000",       "Active-duty personnel"],
        ["Annual fuel spend (Army)",  "$2.7B+ (FY2010)", "70% was for theater operations"],
    ],
    col_widths=[2.5, 1.6, 3.6]
)

h3("What They Need")
body(
    "BCTs operating near power infrastructure — training in CONUS, deployed across Europe (Germany, Italy, Poland), "
    "or staged near allied grid in Korea and Japan — still run diesel generators because autonomous power line "
    "tapping at scale is not currently possible. Chargebotic solves this directly."
)

h3("How They Buy")
bullet("Army SBIR Phase I", bold_part="Army SBIR Phase I", color=C_AMBER)
body("    $250K award, 6-month performance period. The Army has an open RFP specifically for military power solutions innovation. Lowest barrier to entry.", space_after=Pt(2))
bullet("Program Executive Office CS&CSS", bold_part="Program Executive Office CS&CSS", color=C_AMBER)
body("    PEO Combat Support & Combat Service Support manages Army power generation procurement programs.", space_after=Pt(2))
bullet("Other Transaction Authority (OTA)", bold_part="Other Transaction Authority (OTA)", color=C_AMBER)
body("    Faster procurement pathway for innovative technology — bypasses the full FAR acquisition process.", space_after=Pt(2))
bullet("Army Futures Command (AFC)", bold_part="Army Futures Command (AFC)", color=C_AMBER)
body("    The Army's innovation arm. Direct engagement here can accelerate to pilot contracts.", space_after=Pt(2))

# ── 3.2 USSOCOM
h2("3.2 USSOCOM — Highest Priority Customer")

h3("Who They Are")
body(
    "United States Special Operations Command (USSOCOM) is the joint command coordinating Army Special Forces "
    "(Green Berets), Navy SEALs, Rangers, MARSOC, Air Force Special Operations, and Civil Affairs. "
    "They operate in the most austere, denied, and remote environments — often with zero logistics support for days or weeks. "
    "SOCOM pioneered military energy harvesting research (TALOS exoskeleton power system, EHAP assault pack) and has "
    "the most acute version of the energy problem."
)

h3("Key Numbers")
table(
    ["Metric", "Figure", "Notes"],
    [
        ["Total SOF Personnel",       "~73,000",         "Active + reserve + civilian support"],
        ["Countries Active",          "80+",             "Simultaneous operations"],
        ["Acquisition Authority",     "Title 10, §167",  "Can develop and acquire SOF-peculiar equipment independently"],
        ["SBIR Phase I Award",        "$150,000",        "6-month performance period"],
        ["SBIR Phase II Award",       "Up to $1,000,000","2-year performance period"],
        ["Sub-commands",              "5",               "USASOC, NAVSOC, AFSOC, MARSOC, JSOC"],
    ],
    col_widths=[2.5, 1.6, 3.6]
)

h3("What They Need")
body(
    "Small SOCOM teams operating for days or weeks near power infrastructure (roads, villages, allied bases) "
    "currently have no way to tap existing grid energy without Chargebotic. A compact, autonomous system "
    "that silently perches on a power line and delivers continuous energy eliminates the need for fuel "
    "resupply in the most dangerous operational environments."
)

h3("How They Buy")
bullet("SOCOM SBIR", bold_part="SOCOM SBIR", color=C_AMBER)
body("    Phase I = $150K / 6 months. Phase II = up to $1M / 2 years. SOCOM runs its own SBIR program independently.", space_after=Pt(2))
bullet("SOF AT&L (Acquisition, Technology and Logistics)", bold_part="SOF AT&L (Acquisition, Technology and Logistics)", color=C_AMBER)
body("    Streamlined acquisitions office. SOCOM acquisitions move significantly faster than conventional Army.", space_after=Pt(2))
bullet("SOFWERX", bold_part="SOFWERX", color=C_AMBER)
body("    SOCOM's rapid innovation hub in Tampa, FL. Direct engagement with SOF AT&L decision-makers. No prime contractor required.", space_after=Pt(2))

callout(
    "PRIORITY NOTE:  SOCOM should be contacted in parallel with the Army SBIR — not sequentially. "
    "Their independent acquisition authority and faster timelines make them the highest-leverage first engagement "
    "on the government-direct track."
)

# ── 3.3 USMC
h2("3.3 US Marine Corps — Secondary Customer")

h3("Who They Are")
body(
    "The Marine Corps is the US expeditionary force-in-readiness, optimized for rapid deployment from sea to shore. "
    "Marine Expeditionary Units (MEUs) operate with minimal logistics tail and must establish FOB power "
    "quickly before supply chains are established. The USMC has been among the most vocal military branches "
    "about reducing fuel dependency (formal USMC FOB Energy program)."
)

h3("Key Numbers")
table(
    ["Metric", "Figure", "Notes"],
    [
        ["Active Marine Expeditionary Units", "7",         "Primary rapid-deployment operational units"],
        ["Total Active Marines",              "~177,000",  "As of 2024"],
        ["Procurement Office",                "MARCORSYSCOM", "Marine Corps Systems Command"],
        ["Co-procurement",                    "Common",    "Often buys jointly with Army on shared programs"],
    ],
    col_widths=[3.0, 1.6, 3.1]
)

h3("What They Need")
body(
    "MEUs go ashore and establish expeditionary FOBs in coastal and near-urban environments — "
    "which almost always have overhead power line infrastructure nearby. Chargebotic enables "
    "immediate power establishment without waiting for fuel logistics, directly supporting the "
    "Marine expeditionary doctrine of self-sufficiency in the first 72 hours."
)

# ── 3.4 Allied Forces
h2("3.4 Allied Forces and NATO Partners — International Expansion")

body(
    "The US has 32 NATO member allies and key Pacific partners (Japan, South Korea, Australia) "
    "operating on the same doctrine and facing the same energy logistics problem. "
    "A US DoD-validated Chargebotic product accesses this international market through "
    "Foreign Military Sales (FMS) — a US government-managed sales mechanism — or through "
    "direct commercial export with proper ITAR licensing. "
    "The international multiplier is estimated at approximately 3× the US market for a validated product."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — B2B2G STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
h1("4. B2B2G Strategy — Selling Through Defense Tech Primes")

body(
    "In parallel with direct government procurement (SBIR/OTA), Chargebotic can accelerate revenue "
    "and market entry by selling its technology to private defense tech companies that already hold "
    "open government contracts. These companies integrate Chargebotic as a power source add-on and "
    "sell the combined solution to the DoD under their existing contract vehicles. "
    "This is called a Business-to-Business-to-Government (B2B2G) strategy."
)

h2("4.1 Why B2B2G Is the Faster Path")

table(
    ["Factor", "Direct to DoD", "B2B2G Through Primes"],
    [
        ["Time to first contract", "18–36 months",   "3–9 months"],
        ["Procurement complexity", "Full FAR/SBIR",  "Partner handles it"],
        ["Gross margin",           "Full (60–80%)",  "Reduced (integrator takes 20–40%)"],
        ["Revenue risk",           "High (long cycle)", "Lower (partner already funded)"],
        ["DoD credibility",        "Build from scratch", "Inherited from partner"],
        ["IP risk",                "Low",            "Moderate (need NDA + patents first)"],
        ["Recommended approach",   "Run in parallel", "Start here for speed"],
    ],
    col_widths=[2.2, 1.8, 2.8]
)

h2("4.2 B2B2G Target Companies")

h3("1. Chariot Defense  ·  PRIORITY #1")

body(
    "Chariot Defense is the single most important B2B2G target for Chargebotic. "
    "Founded in 2024, Chariot builds Amphora — a modular, software-defined battlefield power distribution system "
    "that senses, prioritizes, and routes electricity across mission-critical systems in real time. "
    "It reduces electronic signatures that would expose a unit's position. "
    "The critical gap: Chariot routes power but has no power source. "
    "Their Amphora system requires electricity to function. Chargebotic is that electricity source."
)

table(
    ["Detail", "Information"],
    [
        ["Founded",          "2024"],
        ["Stealth Exit",     "July 2025"],
        ["Seed Funding",     "$8M led by General Catalyst and XYZ (July 2025)"],
        ["Series A",         "$34M led by Andreessen Horowitz (a16z), February 2026"],
        ["Total Funding",    "$41M"],
        ["Other Investors",  "DCVC, LMNT, Marlinspike, Overmatch, Shield Capital, Ensemble, Trenches Capital"],
        ["Product",          "Amphora — software-defined battlefield power distribution and routing"],
        ["Traction",         "Fielded in force-on-force Army exercises within 6 months of founding"],
        ["Contracts",        "Active sales and contracts with US Army and DIU Project GI"],
        ["Revenue",          "Generating revenue within 12 months of founding"],
        ["HQ",               "San Francisco, CA"],
        ["The pitch",        "You solve distribution. We solve generation. Together you close the loop."],
    ],
    col_widths=[1.8, 5.9]
)

h3("2. Anduril Industries  ·  PRIORITY #2")

body(
    "Anduril is the leading 'new prime' defense tech company, building autonomous drones, counter-drone systems, "
    "border surveillance, missile defense, and the Lattice command-and-control platform. "
    "Their strategy is to own the operating system of the modern battlefield while manufacturing the systems connected to it. "
    "All of Anduril's edge systems — drones, sensors, directed energy weapons — are power-hungry. "
    "Persistent power at the tactical edge is their customers' ceiling."
)

table(
    ["Detail", "Information"],
    [
        ["Valuation",       "$20B+ (2024 funding round)"],
        ["Key products",    "Lattice OS, Roadrunner autonomous drone, Fury autonomous fighter, Ghost-X counter-drone"],
        ["Manufacturing",   "Arsenal-1 — 5M sq ft autonomous weapons factory near Columbus, OH (in development)"],
        ["DoD contracts",   "Multiple; one of the fastest-growing DoD contractors by contract value"],
        ["The pitch",       "Your Lattice platform needs persistent power at the edge. We deliver it from the line above."],
    ],
    col_widths=[1.8, 5.9]
)

h3("3. Shield AI  ·  PRIORITY #3")

body(
    "Shield AI builds Hivemind — autonomy software that allows aircraft and drones to operate without GPS "
    "or communications links (critical for contested electromagnetic environments). "
    "Their V-BAT drone platform and Hivemind stack are deployed operationally. "
    "Shield AI President Brandon Tseng has noted that ~99% of autonomous military systems "
    "cannot function in GPS/comms-jammed environments. Their drones need persistent power to maintain station."
)

table(
    ["Detail", "Information"],
    [
        ["Status",          "Pre-IPO (2026 IPO expected)"],
        ["Key products",    "Hivemind autonomy software, V-BAT drone, Nova autonomous fighter"],
        ["Key claim",       "99% of autonomous military systems fail in comms-jammed environments — Hivemind doesn't"],
        ["DoD contracts",   "Operational deployments; selected for programs involving Anduril-built aircraft"],
        ["The pitch",       "Your drones run out of power. Ours harvests it from the line above them."],
    ],
    col_widths=[1.8, 5.9]
)

h3("4. Mach Industries  ·  PRIORITY #4")

body(
    "Mach Industries manufactures UAVs and loitering munitions at scale for the DoD, "
    "using decentralized manufacturing facilities. Power logistics is their customers' top operational pain. "
    "Chargebotic as a supply chain add-on enables Mach's customers to sustain operations without fuel convoys."
)

table(
    ["Detail", "Information"],
    [
        ["Valuation",       "$1.8B (June 2026) — 4× increase in one year"],
        ["Series B",        "$100M led by Khosla Ventures and Bedrock, with Sequoia"],
        ["Key products",    "UAVs, loitering munitions, decentralized manufacturing"],
        ["Focus",           "Defense manufacturing velocity — making weapons at software speed"],
        ["The pitch",       "Power logistics is your customers' #1 problem. Add Chargebotic to your supply chain."],
    ],
    col_widths=[1.8, 5.9]
)

h3("5. Windlift  ·  WATCH / PARTNER")

body(
    "Windlift is the most structurally similar company to Chargebotic — a tethered drone that generates power "
    "and delivers it via tether to a ground station. The key difference: Windlift harvests wind energy, "
    "not power line energy. This makes their system weather-dependent and unable to operate in low-wind environments. "
    "Windlift has significant DoD backing and has validated the core 'tethered drone as power delivery' concept. "
    "They are either a natural partner or an acqui-hire risk — monitor closely."
)

table(
    ["Detail", "Information"],
    [
        ["DoD Funding",     "$24M+ total DoD support"],
        ["Key contract",    "$30M DoD contract for testing and commercialization"],
        ["Product",         "Autonomous Power Generator (APG) — tethered drone generating electricity from wind"],
        ["System output",   "3 kW rated system per 12-foot-wingspan drone"],
        ["Added function",  "Elevated platform for cameras, sensors, and comms while generating power"],
        ["Risk to Chargebotic", "Same delivery model; could pivot to power line harvesting if they see the market"],
        ["Opportunity",     "License or co-sell agreement; their DoD relationships are valuable"],
    ],
    col_widths=[1.8, 5.9]
)

h2("4.3 B2B2G Risks and Mitigation")

table(
    ["Risk", "Likelihood", "Mitigation"],
    [
        ["Partner copies the technology",      "Medium", "File patents BEFORE any demo. NDA + term sheet immediately."],
        ["Margin compression (20–40%)",        "High",   "Accept for speed; use direct SBIR path in parallel for margin."],
        ["100% dependency on one integrator",  "Medium", "Diversify across 2–3 partners minimum."],
        ["Partner gets acquired, contract ends","Low",   "Build direct DoD relationships through SBIR simultaneously."],
        ["Partner deprioritizes Chargebotic",  "Medium", "Maintain competitive alternatives; never be sole-sourced to one partner."],
    ],
    col_widths=[2.5, 1.2, 3.9]
)

callout(
    "CRITICAL:  File all relevant IP before making contact with any of these companies. "
    "This includes patents on the autonomous perching mechanism, the power harvesting method, "
    "the tether power delivery system, and the ground station interface. "
    "Only then initiate outreach — starting with Chariot Defense."
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Competitive Landscape")

body(
    "No direct competitor exists. No commercial or defense product today offers an autonomous drone "
    "that taps overhead power lines and delivers that energy via tether to a ground station. "
    "The competitive landscape consists of academic research, adjacent energy technologies, and incumbent diesel systems."
)

h2("5.1 Direct Competitors — Power Line Energy Harvesting Drone")

body(
    "None identified. Academic research exists (University of Southern Denmark, 2024: demonstrated autonomous "
    "drone charging from power lines) but this targets drone self-charging, not ground station delivery. "
    "US Patent 7,563,124 covers low-voltage power line harvesting conceptually but has no commercial product. "
    "Chargebotic operates in genuine white space."
)

h2("5.2 Adjacent Technologies (Indirect Competition)")

table(
    ["Company / Technology", "Category", "DoD Funding", "Key Difference", "Threat Level"],
    [
        ["Windlift",              "Tethered wind drone",       "$24M+",    "Wind-dependent, no power line tap, weather-limited",         "Medium — watch"],
        ["DARPA POWER (RTX/Raytheon, Draper, BEAM)", "Laser power relay", "$10M", "Optical/wireless; totally different physics; no tether", "Low"],
        ["Solar + battery hybrids","Renewable field power",   "Widespread","30–50% fuel reduction only; still weather/sun dependent",   "Low"],
        ["Diesel generators (MEP series)","Incumbent power",  "Incumbent", "Fuel-dependent, loud, logistics tail — what we replace",     "None (target)"],
        ["DARPA ExPEDitions",     "Next-gen batteries",       "Active",    "Improves battery density, not power sourcing",               "Low"],
    ],
    col_widths=[2.3, 1.5, 1.3, 2.9, 1.3]
)

h2("5.3 Competitive Positioning Matrix")

table(
    ["Capability",              "Chargebotic", "Windlift", "Solar/Battery", "Diesel Gen.", "DARPA POWER"],
    [
        ["Taps existing infrastructure", "✓", "✗", "✗", "✗", "✗"],
        ["Weather-independent",          "✓", "✗", "✗", "✓", "✓"],
        ["Continuous power (24/7)",      "✓", "✗", "✗", "✓", "Partial"],
        ["Zero fuel required",           "✓", "✓", "✓", "✗", "✓"],
        ["Autonomous deployment",        "✓", "✓", "✗", "✗", "✗"],
        ["No logistics tail",            "✓", "✓", "Partial", "✗", "✓"],
        ["Near-term fielding",           "✓", "✓", "✓", "✓", "✗ (2028+)"],
    ],
    col_widths=[2.5, 1.3, 1.2, 1.4, 1.3, 1.6]
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — KEY PROGRAMS & PROCUREMENT PATH
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Key DoD Programs and Procurement Path")

h2("6.1 Active Programs — Buying Signals")

table(
    ["Program", "Agency", "Size/Status", "Relevance"],
    [
        ["Army SBIR Power Solutions",         "US Army",    "$250K Phase I — open RFP", "Direct entry point for military power innovation"],
        ["SOCOM SBIR",                        "USSOCOM",    "$150K Ph. I / $1M Ph. II", "SOF power and autonomy; own acquisition authority"],
        ["DARPA POWER",                       "DARPA",      "$10M+ funded",             "Signals DoD appetite for novel energy delivery"],
        ["DARPA ExPEDitions",                 "DARPA",      "Active (starts Jan 2027)",  "5–10× battery density — shows DoD wants endurance"],
        ["Army Operational Energy Strategy",  "US Army",    "Formal mandate",            "Top-level DoD mandate creating procurement budget"],
        ["DIU Finance First (Aug 2025)",      "DAF + DIU",  "Active",                    "Energy resiliency initiative for installations"],
        ["Tethered Drones → Army Program",    "US Army",    "Formalizing 2025",          "Tethered drones becoming program-of-record — direct path"],
        ["DoD/GAO Convoy Reports",            "Congress",   "Ongoing",                   "Political cover for procurement; fuel vulnerability cited"],
        ["SOCOM Grand Challenge #3 (TALOS)",  "USSOCOM",    "Historical precedent",      "SOCOM actively funded power innovation for operators"],
    ],
    col_widths=[2.4, 1.3, 1.6, 2.5]
)

h2("6.2 Recommended Entry Sequence")

h3("Step 1 — Secure IP (Before Everything)")
body(
    "File patents on: (1) autonomous perching/attachment mechanism on power lines, "
    "(2) energy harvesting method from overhead lines, (3) tether-based power delivery to ground station, "
    "(4) autonomous identification and navigation to power line targets. "
    "This must happen before any partner or government conversation."
)

h3("Step 2 — Dual Track: B2B2G + Government (Simultaneous)")
bullet("B2B2G: Contact Chariot Defense", bold_part="B2B2G: Contact Chariot Defense", color=C_AMBER)
body("    Chariot needs a power source for Amphora. Under NDA, demonstrate the concept. Pursue a commercial partnership or integration agreement. Fastest path to revenue (3–9 months).", space_after=Pt(2))
bullet("SBIR Track: Submit Army SBIR Phase I ($250K)", bold_part="SBIR Track: Submit Army SBIR Phase I ($250K)", color=C_AMBER)
body("    Army has an open RFP specifically for military power solutions. 6-month performance period. Validates technology with DoD funding.", space_after=Pt(2))
bullet("SBIR Track: Submit SOCOM SBIR Phase I ($150K)", bold_part="SBIR Track: Submit SOCOM SBIR Phase I ($150K)", color=C_AMBER)
body("    Run in parallel with Army SBIR. SOCOM moves faster and has more autonomy. Direct engagement via SOFWERX.", space_after=Pt(2))

h3("Step 3 — OTA Prototype Agreement")
body(
    "After SBIR Phase II ($1M–$1.7M), convert to an Other Transaction Authority (OTA) prototype agreement. "
    "OTAs do not require competitive bidding and allow rapid transition from prototype to production. "
    "Timeline: 18–30 months from SBIR Phase I submission."
)

h3("Step 4 — Program of Record")
body(
    "The long-term goal is to become a budget line item in Army or SOCOM procurement — "
    "a 'program of record' that gets funded annually through the President's Budget Request. "
    "This is the highest-value outcome: predictable, recurring government revenue. Timeline: 3–5 years."
)

h3("Key Offices and People to Engage")
table(
    ["Organization", "Office / Program", "Role"],
    [
        ["US Army",   "PEO CS&CSS",                         "Program Executive Office — buys power generation equipment"],
        ["US Army",   "Army Futures Command (AFC)",          "Innovation arm — accelerates novel technology to pilots"],
        ["USSOCOM",   "SOF AT&L",                           "Acquisition, Technology and Logistics — decision-makers"],
        ["USSOCOM",   "SOFWERX (Tampa, FL)",                 "Innovation hub — direct access without prime contractor"],
        ["DARPA",     "Tactical Technology Office (TTO)",    "Funds novel military hardware programs"],
        ["DoD",       "Defense Innovation Unit (DIU)",       "Bridge between commercial tech and DoD — no SBIR needed"],
        ["USMC",      "MARCORSYSCOM",                        "Marine Corps Systems Command — procurement authority"],
    ],
    col_widths=[1.5, 2.5, 3.8]
)

page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — FULL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Full Summary — The Complete Picture")

h2("7.1 The Opportunity in One Paragraph")

body(
    "The US military spends over $20 billion per year on energy, pays up to $400 per gallon to "
    "deliver fuel to remote forward operating bases, and loses soldiers in fuel resupply convoys. "
    "Diesel generators are the current standard — noisy, fuel-hungry, and logistically expensive. "
    "Existing alternatives (solar, wind, batteries) reduce but do not eliminate the fuel dependency. "
    "Chargebotic is the first system that taps existing overhead power line infrastructure autonomously — "
    "no fuel, no weather dependency, no convoy — and delivers that energy to a ground battery station via tether. "
    "The market is $7.7B–$10.8B today growing at double-digit CAGR. There is no direct competitor. "
    "The fastest path is to call Chariot Defense (already funded at $41M by a16z, routes power but needs a source), "
    "file IP immediately, and run B2B2G and SBIR in parallel."
)

h2("7.2 Complete Data Summary")

table(
    ["Dimension", "Data", "Source / Notes"],
    [
        ["TAM (military power market, 2024)", "$7.7B–$10.8B",      "Mordor Intelligence, IMARC, multiple"],
        ["TAM projection (2030–32)",          "$13.6B–$24.9B",     "8.5–9.7% CAGR"],
        ["Military microgrid sub-market",     "$2.0B (19.1% CAGR)","GM Insights 2024"],
        ["Chargebotic SAM",                   "$400M–$800M",       "Deployable off-grid tactical power"],
        ["DoD annual energy spend",           "$20B+",             "Electric Choice / Army.mil"],
        ["Fully burdened fuel cost (FOB)",    "$400/gallon",       "vs. ~$3 at pump (RAND, Army.mil)"],
        ["Annual fuel spend (Army, FY2010)",  "$2.7B",             "70% theater operations — Army.mil"],
        ["Convoy casualty rate (Afghanistan)","1 per 24 convoys",  "Sustain the Mission Project (DTIC)"],
        ["Fuel-related convoy casualties",    "50%+ (Afg), 30% (Iraq)","2009 data — Army Technology"],
        ["DoD monthly fuel volume in theater","68M gallons",       "2008 average — Army.mil"],
        ["Primary customer",                  "US Army",           "32 active + 27 NG BCTs = 59 total"],
        ["High-priority customer",            "USSOCOM",           "73K SOF, own acquisition authority"],
        ["Secondary customer",                "USMC",              "7 MEUs, expeditionary doctrine"],
        ["BCT personnel",                     "4,000–4,700",       "Per BCT, commanded by Colonel"],
        ["US overseas installations",         "750+ in 80 countries","WarCosts.org 2024"],
        ["Direct competitors",                "0",                 "White space — confirmed"],
        ["B2B2G priority #1",                 "Chariot Defense",   "$41M a16z-led, needs a power source"],
        ["B2B2G priority #2",                 "Anduril",           "$20B+ valuation, Lattice OS"],
        ["Nearest indirect competitor",       "Windlift",          "$24M DoD, wind-based (not power line)"],
        ["Defense tech VC investment (2026)", "$14.6B (Jan–May)",  "Record pace — TechTimes June 2026"],
        ["SBIR Army entry point",             "$250K Phase I",     "Open RFP for power solutions"],
        ["SBIR SOCOM entry point",            "$150K Phase I",     "Faster; own acquisition authority"],
        ["B2B2G time to first contract",      "3–9 months",        "vs. 18–36 months direct to DoD"],
    ],
    col_widths=[2.8, 1.8, 3.1]
)

h2("7.3 Recommended Actions — Priority Order")

bullet("1. File IP now", bold_part="1. File IP now", color=C_RED)
body("    Patent the perching mechanism, energy harvesting method, tether delivery system, and autonomous line identification. Do this before any external conversation.", space_after=Pt(2))

bullet("2. Contact Chariot Defense under NDA", bold_part="2. Contact Chariot Defense under NDA", color=C_AMBER)
body("    They are the single most natural partner — $41M funded, a16z-backed, Army contracts active, and they have no power source for their Amphora distribution system. Fastest path to revenue.", space_after=Pt(2))

bullet("3. Submit Army SBIR Phase I ($250K)", bold_part="3. Submit Army SBIR Phase I ($250K)", color=C_AMBER)
body("    Army has an open RFP for military power solutions. Run this in parallel with Chariot conversation.", space_after=Pt(2))

bullet("4. Submit SOCOM SBIR Phase I ($150K)", bold_part="4. Submit SOCOM SBIR Phase I ($150K)", color=C_AMBER)
body("    Run simultaneously with Army SBIR. Engage SOFWERX directly in Tampa for accelerated access.", space_after=Pt(2))

bullet("5. Engage DIU (Defense Innovation Unit)", bold_part="5. Engage DIU (Defense Innovation Unit)", color=C_BLUE)
body("    DIU bridges commercial tech and DoD without SBIR process. Good parallel track for rapid pilot.", space_after=Pt(2))

bullet("6. Expand B2B2G to Anduril and Shield AI", bold_part="6. Expand B2B2G to Anduril and Shield AI", color=C_BLUE)
body("    After Chariot partnership is secured, approach Anduril and Shield AI as the next integration targets.", space_after=Pt(2))

callout(
    "BOTTOM LINE:  The DoD has a $20B+ annual energy problem with no current solution for autonomous power line tapping. "
    "Chariot Defense just raised $41M to distribute power — but they have no source. "
    "Chargebotic is the source. File IP today. Call Chariot tomorrow. "
    "Run SBIR in parallel. The white space is real and the window is open."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_border_bottom(p, color="D4810F", size=4)
run = p.add_run(
    "Chargebotic  ·  anis@chargebotic.com  ·  chargebotic.com  ·  Confidential — June 2026"
)
run.font.size = Pt(9)
run.font.color.rgb = C_GRAY
run.font.name = "Calibri"
run.italic = True


# ── Save ──────────────────────────────────────────────────────────────────────
output = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "defense-market-research.docx"
)
doc.save(output)
print(f"Saved: {output}")
